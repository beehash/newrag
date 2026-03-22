#!/usr/bin/env python3
"""
将 docs 目录中的所有文件入库到 milvus 数据库和 elasticsearch 数据库
"""

import os
import sys
import uuid
import time
from datetime import datetime
import io
import pdfplumber
from docx import Document

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 导入配置
from config import MILVUS_CONFIG, FILE_TYPES

# 导入必要的模块
from ingestion.embedding import EmbeddingService
from ingestion.mv_store import MVStore
from ingestion.es_store import ESStore
from ingestion.chunker import chunk_text
import asyncio

def parse_pdf_content(file_content: bytes) -> str:
    """解析PDF文件内容"""
    content = []
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
    except Exception as e:
        print(f"PDF解析警告: {str(e)}")
    return "\n".join(content)

def parse_docx_content(file_content: bytes) -> str:
    """解析DOCX文件内容"""
    content = []
    try:
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            if paragraph.text:
                content.append(paragraph.text)
    except Exception as e:
        print(f"DOCX解析警告: {str(e)}")
    return "\n".join(content)

def parse_file_content(file_content: bytes, file_ext: str) -> str:
    """直接解析文件内容"""
    try:
        if file_ext == ".pdf":
            return parse_pdf_content(file_content)
        elif file_ext == ".docx":
            return parse_docx_content(file_content)
        elif file_ext in [".txt", ".md"]:
            return file_content.decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        raise Exception(f"文件解析失败: {str(e)}")

async def process_file(file_path, embedding_service, mv_store, es_store):
    """
    处理单个文件并入库
    
    Args:
        file_path: 文件路径
        embedding_service: Embedding服务实例
        mv_store: MVStore实例
        es_store: ESStore实例
        
    Returns:
        bool: 处理是否成功
    """
    doc_id = None
    try:
        print(f"\n{'='*60}")
        print(f"处理文件: {file_path}")
        print(f"{'='*60}")
        
        # 获取文件名和扩展名
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        # 检查文件类型
        if file_ext not in FILE_TYPES:
            print(f"跳过不支持的文件类型: {file_ext}")
            return False
        
        # 生成doc_id
        doc_id = str(uuid.uuid4())
        print(f"生成的文档ID: {doc_id}")
        
        # 读取文件内容
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
            print(f"文件读取成功，大小: {len(file_content)} 字节")
        except Exception as e:
            print(f"文件读取失败: {str(e)}")
            return False
        
        # 解析文件内容
        content = parse_file_content(file_content, file_ext)
        if not content:
            print(f"文件内容为空，跳过处理")
            return False
        print(f"文件解析成功，内容长度: {len(content)} 字符")
        
        # 对文本进行切块
        chunks = chunk_text(content)
        if not chunks:
            print(f"文本切块结果为空，跳过处理")
            return False
        print(f"文本切块完成，共 {len(chunks)} 个块")
        
        # 为每个块生成embedding
        chunk_embeddings = []
        try:
            for i, chunk in enumerate(chunks):
                print(f"正在处理第 {i+1}/{len(chunks)} 个块，块长度: {len(chunk)}")
                embedding = embedding_service.embed_text(chunk)
                chunk_embeddings.append({
                    "chunk": chunk,
                    "embedding": embedding
                })
            print("所有块的embedding生成完成")
        except Exception as e:
            print(f"Embedding生成失败: {str(e)}")
            return False
        
        # 创建对象列表为数据入库做准备
        documents_for_db = []
        for i, item in enumerate(chunk_embeddings):
            # 生成title：filename去掉后缀
            title = filename.rsplit('.', 1)[0] if '.' in filename else filename
            
            documents_for_db.append({
                "filename": filename,
                "doc_id": doc_id,
                "text": item["chunk"],
                "vector": item["embedding"],
                "chunk_id": i + 1,
                "type": FILE_TYPES[file_ext],
                "title": title,
                "create_at": datetime.now().isoformat()
            })
        
        # 步骤1: 插入数据到Milvus数据库
        print(f"\n步骤1: 插入数据到Milvus数据库...")
        try:
            mv_store.insert(documents_for_db)
            print(f"✓ 文件 {filename} 成功插入到Milvus数据库，共 {len(documents_for_db)} 条记录")
        except Exception as e:
            print(f"✗ 插入Milvus数据库失败: {str(e)}")
            return False
        
        # 步骤2: 插入数据到Elasticsearch数据库
        print(f"\n步骤2: 插入数据到Elasticsearch数据库...")
        try:
            # 准备Elasticsearch数据（包含所有需要的字段）
            es_documents = []
            for doc in documents_for_db:
                es_documents.append({
                    "chunk_id": doc["chunk_id"],
                    "doc_id": doc["doc_id"],
                    "text": doc["text"],
                    "filename": doc["filename"],
                    "title": doc["title"],
                    "create_at": doc["create_at"],
                    "type": doc["type"]
                })
            
            # 异步插入数据
            await es_store.add(es_documents)
            print(f"✓ 文件 {filename} 成功插入到Elasticsearch数据库，共 {len(es_documents)} 条记录")
            print(f"\n{'='*60}")
            print(f"✓ 文件 {filename} 处理完成！")
            print(f"{'='*60}")
            return True
        except Exception as e:
            print(f"✗ 插入Elasticsearch数据库失败: {str(e)}")
            # 如果Elasticsearch插入失败，删除已插入到Milvus的数据
            print(f"\n回滚: 删除Milvus中的数据，doc_id: {doc_id}")
            try:
                mv_store.deleteDocument(doc_id)
                print(f"✓ 已成功回滚Milvus中的数据")
            except Exception as rollback_error:
                print(f"✗ 回滚Milvus数据失败: {str(rollback_error)}")
            return False
            
    except Exception as e:
        print(f"✗ 处理文件 {file_path} 时出错: {str(e)}")
        # 如果在处理过程中出错，尝试清理已插入的数据
        if doc_id:
            try:
                mv_store.deleteDocument(doc_id)
                print(f"✓ 已清理Milvus中的数据，doc_id: {doc_id}")
            except Exception as cleanup_error:
                print(f"✗ 清理Milvus数据失败: {str(cleanup_error)}")
        return False

async def main():
    """
    主函数
    """
    # 文档目录路径
    docs_dir = "/Users/eanaminsey/Documents/graduate-projects/agents/docs"
    
    print(f"\n{'#'*60}")
    print(f"# 批量入库脚本")
    print(f"# 目标目录: {docs_dir}")
    print(f"# 数据库: Milvus + Elasticsearch")
    print(f"{'#'*60}\n")
    
    # 检查目录是否存在
    if not os.path.exists(docs_dir):
        print(f"✗ 目录不存在: {docs_dir}")
        return
    
    # 初始化Embedding服务
    print("初始化Embedding服务...")
    try:
        embedding_service = EmbeddingService()
        print("✓ Embedding服务初始化成功！")
    except Exception as e:
        print(f"✗ Embedding服务初始化失败: {str(e)}")
        return
    
    # 初始化MVStore（使用zw_docs集合）
    print("\n初始化Milvus数据库...")
    try:
        mv_store = MVStore(
            collection_name="zw_docs"
        )
        
        # 检查连接状态
        if mv_store.client is None:
            print("✗ Milvus数据库连接失败，退出脚本")
            return
        
        print("✓ Milvus数据库连接成功！")
    except Exception as e:
        print(f"✗ Milvus数据库连接失败: {str(e)}")
        return
    
    # 初始化ESStore
    print("\n初始化Elasticsearch数据库...")
    try:
        es_store = ESStore()
        
        # 异步连接Elasticsearch
        await es_store._connect()
        
        if es_store.client is None:
            print("✗ Elasticsearch数据库连接失败，退出脚本")
            return
        
        print("✓ Elasticsearch数据库连接成功！")
    except Exception as e:
        print(f"✗ Elasticsearch数据库连接失败: {str(e)}")
        return
    
    # 遍历目录中的所有文件
    print(f"\n开始扫描目录: {docs_dir}")
    
    # 支持的文件扩展名
    supported_exts = list(FILE_TYPES.keys())
    
    # 收集所有支持的文件
    files_to_process = []
    for root, dirs, files in os.walk(docs_dir):
        for file in files:
            file_ext = os.path.splitext(file)[1].lower()
            if file_ext in supported_exts:
                files_to_process.append(os.path.join(root, file))
    
    print(f"✓ 发现 {len(files_to_process)} 个支持的文件")
    
    if len(files_to_process) == 0:
        print("\n没有需要处理的文件，脚本结束")
        return
    
    # 处理每个文件
    success_count = 0
    failure_count = 0
    start_time = time.time()
    
    for i, file_path in enumerate(files_to_process, 1):
        print(f"\n进度: [{i}/{len(files_to_process)}]")
        
        if await process_file(file_path, embedding_service, mv_store, es_store):
            success_count += 1
        else:
            failure_count += 1
        
        # 每处理3个文件，暂停2秒，避免请求过于频繁
        if i % 3 == 0 and i < len(files_to_process):
            print(f"\n暂停2秒，避免请求过于频繁...")
            time.sleep(2)
    
    # 完成统计
    end_time = time.time()
    total_time = end_time - start_time
    
    print(f"\n{'#'*60}")
    print(f"# 处理完成统计")
    print(f"{'#'*60}")
    print(f"✓ 成功: {success_count} 个文件")
    print(f"✗ 失败: {failure_count} 个文件")
    print(f"总计: {success_count + failure_count} 个文件")
    print(f"总耗时: {total_time:.2f} 秒")
    print(f"平均耗时: {total_time/(success_count + failure_count):.2f} 秒/文件")
    print(f"{'#'*60}\n")
    
    # 关闭连接
    print("关闭数据库连接...")
    try:
        await es_store.close()
        print("✓ Elasticsearch连接已关闭")
    except Exception as e:
        print(f"✗ 关闭Elasticsearch连接失败: {str(e)}")
    
    print("\n✓ 脚本执行完成！")

if __name__ == "__main__":
    asyncio.run(main())
