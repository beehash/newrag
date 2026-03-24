import sys
import os
import uuid
import threading
import time
from datetime import datetime
import tempfile
import pdfplumber
from docx import Document
import uvicorn
import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入配置
from config import MILVUS_CONFIG, LLM_CONFIG, ZHIPU_CONFIG, SERVER_CONFIG, FILE_TYPES, SEARCH_CONFIG

from fastapi import FastAPI, UploadFile, File, HTTPException, Form, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from contextlib import asynccontextmanager
import json
from ingestion.chunker import chunk_text
from ingestion.embedding import EmbeddingService
from ingestion.mv_store import MVStore
from ingestion.es_store import ESStore
from pipline.query_rewrite import QueryRewriter
from pipline.rerank import Reranker
# from pipline.context_compress import ContextCompressor
from retrieval.hybrid_search import HybridSearch
from llm.client import LLMClient

# 生命周期事件处理器
@asynccontextmanager
async def lifespan(app):
    """
    应用生命周期管理
    """
    # 启动时的初始化工作
    yield
    # 关闭时的清理工作
    global llm_client
    if llm_client:
        await llm_client.close()
        print("LLM客户端已关闭")

app = FastAPI(lifespan=lifespan)

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 在生产环境中应该设置具体的前端域名
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 创建API路由
api_router = APIRouter(prefix="/api")

# 添加一个简单的测试接口
@api_router.get("/test")
async def test():
    """测试接口"""
    return {"message": "API测试成功", "status": "ok"}

# 全局MVStore实例
mv_store = None

# 全局ESStore实例
es_store = None

# 全局EmbeddingService实例
embedding_service = None

# 全局查询改写器实例
rewriter = None

# 全局重排序器实例
reranker = None

# 全局上下文压缩器实例
# compressor = None

# 全局LLM客户端实例
llm_client = None

# BM25索引缓存
bm25_cache = {
    "documents": None,
    "documents_for_bm25": None,
    "last_update": 0
}

# 初始化数据库连接
def init_db():
    """
    初始化数据库连接
    """
    global mv_store, embedding_service, es_store
    print("正在连接Milvus数据库...")
    try:
        # 连接数据库
        host = MILVUS_CONFIG.get('host', '127.0.0.1')
        port = MILVUS_CONFIG.get('port', '19530')
        print(f"连接到: {host}:{port}")
        print(f"使用数据库: {MILVUS_CONFIG['db_name']}")
        
        mv_store = MVStore(
            connection_name=MILVUS_CONFIG["connection_name"],
            uri=f"tcp://{host}:{port}",
            token=MILVUS_CONFIG["token"],
            db_name=MILVUS_CONFIG["db_name"]
        )
        
        # 检查客户端是否成功初始化
        if mv_store.client is None:
            raise Exception("Milvus客户端初始化失败")
        
        print("Milvus数据库连接成功！")
    except Exception as e:
        print(f"Milvus数据库连接失败: {str(e)}")
        print("服务将在无数据库连接的情况下启动，部分功能可能不可用")
        mv_store = None
    
    # 初始化Elasticsearch连接
    print("正在连接Elasticsearch数据库...")
    try:
        es_store = ESStore()
        print("Elasticsearch数据库连接成功！")
    except Exception as e:
        print(f"Elasticsearch数据库连接失败: {str(e)}")
        print("服务将在无Elasticsearch连接的情况下启动，部分功能可能不可用")
        es_store = None
    
    # 初始化Embedding服务（在后台线程中运行）
    print("正在初始化Embedding服务...")
# 初始化数据库连接
init_db()

# 初始化Embedding服务
def init_embedding_service():
    global embedding_service
    try:
        embedding_service = EmbeddingService()
        print("Embedding服务初始化成功！")
    except Exception as e:
        print(f"Embedding服务初始化失败: {str(e)}")
        print("服务将在无Embedding服务的情况下启动，部分功能可能不可用")

# 预加载所有模型和组件
def init_components():
    """
    预加载所有模型和组件
    """
    global rewriter, reranker, llm_client
    
    print("正在预加载模型和组件...")
    
    # 初始化查询改写器
    try:
        rewriter = QueryRewriter(
            api_key=ZHIPU_CONFIG["api_key"],
            base_url=ZHIPU_CONFIG["base_url"]
        )
        print("查询改写器初始化成功！")
    except Exception as e:
        print(f"查询改写器初始化失败: {str(e)}")
    
    # 初始化重排序器
    try:
        reranker = Reranker()
        print("重排序器初始化成功！")
    except Exception as e:
        print(f"重排序器初始化失败: {str(e)}")
    
    # # 初始化上下文压缩器
    # try:
    #     compressor = ContextCompressor()
    #     print("上下文压缩器初始化成功！")
    # except Exception as e:
    #     print(f"上下文压缩器初始化失败: {str(e)}")
    
    # 初始化LLM客户端
    try:
        llm_client = LLMClient(
            api_key=LLM_CONFIG["api_key"],
            base_url=LLM_CONFIG["base_url"]
        )
        print("LLM客户端初始化成功！")
    except Exception as e:
        print(f"LLM客户端初始化失败: {str(e)}")
    
    print("模型和组件预加载完成！")

# 同步初始化所有服务
print("正在初始化Embedding服务...")
init_embedding_service()
print("正在预加载模型和组件...")
init_components()

# 上传文件保存目录
UPLOAD_DIR = SERVER_CONFIG["upload_dir"]
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# 支持的文件类型
SUPPORTED_FILE_TYPES = FILE_TYPES

@api_router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """文件上传接口"""
    print("文件上传接口被调用")
    print(f"接收到的文件名: {file.filename}")
    
    # 检查文件类型
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in SUPPORTED_FILE_TYPES:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型，仅支持: {', '.join(SUPPORTED_FILE_TYPES.keys())}")
    
    # 生成doc_id
    doc_id = str(uuid.uuid4())
    print(f"生成的文档ID: {doc_id}")
    
    # 直接读取文件内容
    try:
        file_content = await file.read()
        print(f"文件读取成功，大小: {len(file_content)} 字节")
    except Exception as e:
        print(f"文件读取失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件读取失败: {str(e)}")
    
    # 解析文件内容
    content = parse_file_content(file_content, file_ext)
    print(f"文件解析成功，内容长度: {len(content)} 字符")
    
    # 对文本进行切块
    chunks = chunk_text(content)
    print(f"文本切块完成，共 {len(chunks)} 个块")
    
    # 等待Embedding服务初始化完成
    max_wait_time = 30  # 最多等待30秒
    wait_time = 0
    while embedding_service is None and wait_time < max_wait_time:
        print(f"等待Embedding服务初始化... ({wait_time}/{max_wait_time}秒)")
        time.sleep(1)
        wait_time += 1
    
    # 检查Embedding服务是否可用
    if embedding_service is None:
        print("Embedding服务未初始化")
        raise HTTPException(status_code=500, detail="Embedding服务未初始化")
    
    # 为每个块生成embedding
    chunk_embeddings = []
    try:
        for i, chunk in enumerate(chunks):
            print(f"正在处理第 {i+1}/{len(chunks)} 个块")
            embedding = embedding_service.embed_text(chunk)
            chunk_embeddings.append({
                "chunk": chunk,
                "embedding": embedding
            })
        print("所有块的embedding生成完成")
    except Exception as e:
        print(f"Embedding生成失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Embedding生成失败: {str(e)}")
    
    # 创建对象列表为数据入库做准备
    documents_for_db = []
    for i, item in enumerate(chunk_embeddings):
        # 生成title：filename去掉后缀
        title = file.filename.rsplit('.', 1)[0] if '.' in file.filename else file.filename
        
        documents_for_db.append({
            "filename": file.filename,
            "doc_id": doc_id,
            "text": item["chunk"],
            "vector": item["embedding"],
            "chunk_id": i + 1,
            "type": SUPPORTED_FILE_TYPES[file_ext],
            "title": title,
            "create_at": datetime.now().isoformat()
        })
    
    # 尝试插入数据到Milvus数据库
    if mv_store is not None:
        try:
            mv_store.insert(documents_for_db)
            print(f"文件 {file.filename} 成功插入到Milvus数据库")
        except Exception as e:
            print(f"插入Milvus数据库失败: {str(e)}")
    else:
        print("Milvus数据库连接未初始化，跳过数据库插入")
    
    # 尝试插入数据到Elasticsearch数据库
    if es_store is not None:
        try:
            # 准备Elasticsearch数据
            es_documents = []
            for doc in documents_for_db:
                es_documents.append({
                    "chunk_id": doc["chunk_id"],
                    "doc_id": doc["doc_id"],
                    "text": doc["text"],
                    "filename": doc["filename"],
                    "title": doc["title"],
                    "create_at": doc["create_at"],
                    "type": doc["type"]
                })
            
            # 异步插入数据
            await es_store.add(es_documents)
            print(f"文件 {file.filename} 成功插入到Elasticsearch数据库，共 {len(es_documents)} 条记录")
        except Exception as e:
            print(f"插入Elasticsearch数据库失败: {str(e)}")
            # 如果Elasticsearch插入失败，删除已插入到Milvus的数据
            if mv_store is not None:
                print(f"回滚: 删除Milvus中的数据，doc_id: {doc_id}")
                try:
                    mv_store.deleteDocument(doc_id)
                    print("已成功回滚Milvus中的数据")
                except Exception as rollback_error:
                    print(f"回滚Milvus数据失败: {str(rollback_error)}")
    else:
        print("Elasticsearch数据库连接未初始化，跳过数据库插入")
    
    # 返回结果
    return {
        "filename": file.filename,
        "title": title,
        "content": content[:1000] + "..." if len(content) > 1000 else content,  # 只返回前1000个字符
        "size": len(content),
        "type": SUPPORTED_FILE_TYPES[file_ext],
        "chunks": chunks,
        "chunk_count": len(chunks),
        "doc_id": doc_id
    }

def parse_file(file_path: str, file_ext: str) -> str:
    """解析不同类型的文件"""
    try:
        if file_ext == ".pdf":
            return parse_pdf(file_path)
        elif file_ext == ".docx":
            return parse_docx(file_path)
        elif file_ext in [".txt", ".md"]:
            return parse_txt(file_path)
        else:
            return ""
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

def parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.append(text)
    return "\n".join(content)

def parse_docx(file_path: str) -> str:
    """解析DOCX文件"""
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            content.append(paragraph.text)
    return "\n".join(content)

def parse_txt(file_path: str) -> str:
    """解析TXT和MD文件"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_file_content(file_content: bytes, file_ext: str) -> str:
    """直接解析文件内容"""
    try:
        if file_ext == ".pdf":
            return parse_pdf_content(file_content)
        elif file_ext == ".docx":
            return parse_docx_content(file_content)
        elif file_ext in [".txt", ".md"]:
            return file_content.decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

def parse_pdf_content(file_content: bytes) -> str:
    """解析PDF文件内容"""
    import io
    content = []
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.append(text)
    return "\n".join(content)

def parse_docx_content(file_content: bytes) -> str:
    """解析DOCX文件内容"""
    import io
    doc = Document(io.BytesIO(file_content))
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            content.append(paragraph.text)
    return "\n".join(content)

@api_router.get("/docs")
async def get_docs():
    """获取所有文档列表"""
    # 使用全局MVStore实例
    global mv_store
    if mv_store is None:
        print("数据库连接未初始化，返回空列表")
        return []
    try:
        return mv_store.getAllDocuments()
    except Exception as e:
        print(f"获取文档列表失败: {str(e)}")
        return []

@api_router.get("/doc/{doc_id}")
async def get_doc(doc_id: str):
    """获取单个文档内容"""
    # 使用全局MVStore实例
    global mv_store
    if mv_store is None:
        print("数据库连接未初始化，返回空内容")
        return {"content": ""}
    try:
        doc = mv_store.get_document(doc_id)
        return {"content": doc["content"] if doc else ""}
    except Exception as e:
        print(f"获取文档失败: {str(e)}")
        return {"content": ""}

@api_router.delete("/doc/{doc_id}")
async def delete_doc(doc_id: str):
    """删除文档"""
    # 使用全局MVStore和ESStore实例
    global mv_store, es_store
    if mv_store is None:
        print("数据库连接未初始化，返回失败")
        return {"success": False}
    try:
        # 删除Milvus中的数据
        success = mv_store.deleteDocument(doc_id)
        if not success:
            print("删除Milvus数据失败")
            return {"success": False}
        
        # 删除Elasticsearch中的数据
        if es_store is not None:
            try:
                await es_store.delete_by_doc_id(doc_id)
                print(f"成功删除Elasticsearch中doc_id为 {doc_id} 的数据")
            except Exception as e:
                print(f"删除Elasticsearch数据失败: {str(e)}")
                # 即使Elasticsearch删除失败，也返回成功，因为Milvus已经删除了
        
        return {"success": True}
    except Exception as e:
        print(f"删除文档失败: {str(e)}")
        return {"success": False}

@api_router.post("/query")
async def query(request: dict):
    """查询接口"""
    import time
    start_time = time.time()
    
    try:
        # 获取查询参数
        q = request.get("q", "").strip()
        topk = request.get("topk", 5)
        
        # 验证参数
        if not q:
            raise HTTPException(status_code=400, detail="查询内容不能为空")
        
        # 限制topk范围
        topk = min(max(topk, 1), 10)
        
        # 使用全局MVStore实例
        global mv_store, rewriter, reranker, llm_client
        if mv_store is None:
            print("数据库连接未初始化，尝试重新连接...")
            init_db()
        if mv_store is None:
            raise HTTPException(status_code=500, detail="数据库连接失败")
        
        # 获取embedding模型
        embedding_service = mv_store.getEmbeddingModel()
        
        # 步骤1: 执行查询改写和意图识别
        step1_start = time.time()
        if rewriter is None:
            rewriter = QueryRewriter(
                api_key=ZHIPU_CONFIG["api_key"],
                base_url=ZHIPU_CONFIG["base_url"]
            )
        
        # 执行综合查询分析（一次调用完成意图识别和查询改写）
        analysis_result = rewriter.analyze_query(q, model=ZHIPU_CONFIG["intent_model"], num_rewrites=2)
        intent_result = analysis_result.get("intent", {})
        rewritten_queries = analysis_result.get("rewritten_queries", [])
        print(f"意图识别结果: {intent_result}")
        print(f"改写查询: {rewritten_queries}")
        
        # 构建查询列表：原始查询 + 改写查询
        all_queries = intent_result.get('entities', []) + [q]+ rewritten_queries
        all_queries = list(set(all_queries))  # 去重
        # all_queries = ';'.join(all_queries)
        
        print(f"步骤1完成（查询改写+意图识别）: {time.time() - step1_start:.2f}s  {all_queries}")
        
        # 步骤2: 执行混合搜索
        step2_start = time.time()
        hybrid_search = HybridSearch(mv_store)
        # 只为原始查询生成embedding，减少计算
        print(f"查询query embedding: {q + ' ' + ' '.join(intent_result.get('entities', []))}")
        # 从意图识别结果中获取实体词
        entities = intent_result.get('entities', [])
        print(f"获取到的实体词: {entities}")
        
        # 执行混合搜索
        all_results = []
        for query_text in all_queries:
            query_embedding = embedding_service.embed_text(query_text + ' ' + ' '.join(intent_result.get('entities', [])))
            # 对于所有查询，使用混合搜索，传递实体词用于boost
            search_results = await hybrid_search.search(query_text, query_embedding, limit=10, entities=entities)
            all_results.extend(search_results["vector_results"])
            all_results.extend(search_results["es_results"])
            
        print(f"步骤2完成（混合搜索）: {time.time() - step2_start:.2f}s")
        # 3. 融合搜索结果（使用RRF算法）
        rrf_results = hybrid_search._rrf_fusion(all_results, k=60, limit=topk)
        
        # 步骤3: 处理搜索结果，只保留需要的字段
        step3_start = time.time()
        resources = []
        for result in rrf_results:
            # 只保留需要的字段
            resource = {
                "doc_id": result.get("doc_id"),
                "chunk_id": result.get("chunk_id"),
                "title": result.get("title"),
                "filename": result.get("filename"),
                "create_at": result.get("create_at"),
                "similarity": result.get("similarity"),
                "_similarity": result.get("_similarity"),
                "score": result.get("score", 0)
            }
            
            resources.append(resource)
        
        # 按相似度排序并限制数量
        resources.sort(key=lambda x: x["score"], reverse=True)
        # resources = resources[:topk]
        print(f"步骤3完成（处理搜索结果）: {time.time() - step3_start:.2f}s")
        
        # 步骤4: 重排序（使用预加载的reranker，且只对前10个重排序）
        step4_start = time.time()
        if reranker is None:
            reranker = Reranker()
        # 只对前10个资源重排序（如果资源少于10个则全部重排序）
        rerank_candidates = resources[:min(10, len(resources))]
        reranked_resources = reranker.rerank(q, rerank_candidates, top_k=topk)
        print(f"步骤4完成（重排序）: {time.time() - step4_start:.2f}s")
        
        # # 步骤5: 上下文压缩（使用预加载的compressor）
        # step5_start = time.time()
        # if compressor is None:
        #     compressor = ContextCompressor()
        # # 压缩上下文，减少大模型输入长度
        # compressed_resources = compressor.compress(q, reranked_resources, top_k=topk)
        # print(f"步骤5完成（上下文压缩）: {time.time() - step5_start:.2f}s")
        
        # 步骤6: 构建context
        step6_start = time.time()
        context = ""
        for i, doc in enumerate(reranked_resources):
            context += f"来源 {i+1}: {doc.get('text', '')}\n\n"
        print(f"步骤6完成（构建context）: {time.time() - step6_start:.2f}s")
        
        # 步骤7: 初始化LLM客户端（使用预加载的）
        step7_start = time.time()
        if llm_client is None:
            llm_client = LLMClient(
                api_key=LLM_CONFIG["api_key"],
                base_url=LLM_CONFIG["base_url"]
            )
        print(f"步骤7完成（初始化LLM）: {time.time() - step7_start:.2f}s")
        
        total_time = time.time() - start_time
        print(f"总耗时: {total_time:.2f}s")
        
        # 定义流式生成器
        async def generate():
            # 先发送初始化数据（意图和资源）
            init_data = {
                "type": "init",
                "intent": intent_result,
                "resources": reranked_resources
            }
            yield f"data: {json.dumps(init_data)}\n\n"
            
            # 异步流式获取大模型回复
            async for chunk in llm_client.generate_response_stream(context, q):
                chunk_data = {
                    "type": "chunk",
                    "content": chunk
                }
                yield f"data: {json.dumps(chunk_data)}\n\n"
            
            # 发送结束信号
            end_data = {"type": "end"}
            yield f"data: {json.dumps(end_data)}\n\n"
        
        return StreamingResponse(generate(), media_type="text/event-stream")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")

@api_router.post("/query_rewrite")
async def query_rewrite(request: dict):
    """查询改写接口"""
    try:
        # 获取查询参数
        q = request.get("q", "").strip()
        
        # 验证参数
        if not q:
            raise HTTPException(status_code=400, detail="查询内容不能为空")
        
        # 初始化查询改写器
        rewriter = QueryRewriter(
            api_key=ZHIPU_CONFIG["api_key"],
            base_url=ZHIPU_CONFIG["base_url"]
        )
        
        # 生成改写查询
        rewritten_queries = rewriter.rewrite_query(q)
        
        return {
            "original_query": q,
            "rewritten_queries": rewritten_queries
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询改写失败: {str(e)}")

# 挂载API路由
app.include_router(api_router)



if __name__ == "__main__":
    uvicorn.run("api.server:app", host=SERVER_CONFIG["host"], port=SERVER_CONFIG["port"], reload=True)
